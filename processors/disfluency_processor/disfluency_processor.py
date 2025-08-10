import os
import sys
import re
import torch
import logging
from pathlib import Path
from typing import List, Tuple, Generator, Optional, Union
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Add the joint-disfluency-detector-and-parser/src directory to the path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root / "joint-disfluency-detector-and-parser"))
sys.path.insert(0, str(project_root / "joint-disfluency-detector-and-parser" / "src"))

# Now import the required modules
from parse_nk import NKChartParser
import parse_nk
from trees import InternalTreebankNode, LeafTreebankNode

class DisfluencyProcessor:
    def __init__(self, model_path: str):
        """Initialize the disfluency processor with the specified model.
        
        Args:
            model_path: Path to the trained model file (.pt)
        """
        # Force using CPU to avoid MPS compatibility issues
        self.device = torch.device("cpu")
        print("Using CPU (MPS disabled due to compatibility issues)")  # (no GPU acceleration available)")
            
        # Load the model
        self.parser = self._load_model(model_path)
        
        # Set model to evaluation mode
        self.parser.eval()
        
        # Get dummy tag for parsing
        if hasattr(self.parser, 'tag_vocab'):
            if 'UNK' in self.parser.tag_vocab.indices:
                self.dummy_tag = 'UNK'
            else:
                self.dummy_tag = self.parser.tag_vocab.value(0)
        else:
            self.dummy_tag = 'UNK'
    
    def _load_model(self, model_path: str) -> NKChartParser:
        """Load the pre-trained disfluency detection model.
        
        Args:
            model_path: Path to the model file
            
        Returns:
            Loaded NKChartParser model
        """
        print(f"Loading model from {model_path}...")
        
        # Change to the joint-disfluency-detector-and-parser directory where BERT files are located
        original_cwd = os.getcwd()
        model_dir = Path(model_path).parent.parent / 'model'
        if model_dir.exists():
            os.chdir(str(model_dir))
            print(f"Changed working directory to: {os.getcwd()}")
        
        try:
            # Load model info (on CPU first) with weights_only=False for compatibility
            # Note: Setting weights_only=False is required for this model but be cautious with untrusted files
            info = torch.load(model_path, map_location='cpu', weights_only=False)
            
            # Initialize model using from_spec (handles device placement internally)
            parser = NKChartParser.from_spec(info['spec'], info['state_dict'])
            
            # Move model to the appropriate device
            parser = parser.to(self.device)
            print(f"Model loaded on device: {next(parser.parameters()).device}")
            
            return parser
        finally:
            # Restore original working directory
            os.chdir(original_cwd)
    
    def _is_disfluent_node(self, node: InternalTreebankNode) -> bool:
        """Check if a node is a disfluency node (EDITED, PRN, UH, etc.)"""
        if isinstance(node, LeafTreebankNode):
            return False
        label = node.label.upper()
        return any(disfluency in label for disfluency in ['EDITED', 'PRN', 'UH', 'INTJ'])

    def _clean_tree(self, node: InternalTreebankNode) -> Optional[InternalTreebankNode]:
        """
        Recursively clean a parse tree by removing disfluent nodes and their children.
        
        Args:
            node: The root of the (sub)tree to clean
            
        Returns:
            The cleaned tree, or None if the entire subtree should be removed
        """
        if isinstance(node, LeafTreebankNode):
            return node
            
        # Process children first
        cleaned_children = []
        for child in node.children:
            cleaned_child = self._clean_tree(child)
            if cleaned_child is not None:
                cleaned_children.append(cleaned_child)
                
        # If this is a disfluency node, return None to remove it
        if self._is_disfluent_node(node):
            return None
            
        # If all children were removed, return None
        if not cleaned_children and not isinstance(node, LeafTreebankNode):
            return None
            
        # Create a new node with cleaned children
        return InternalTreebankNode(node.label, cleaned_children)

    def _extract_clean_text(self, node: InternalTreebankNode) -> str:
        """Extract clean text from a tree node, skipping disfluent segments."""
        if isinstance(node, LeafTreebankNode):
            return node.word
            
        # Skip disfluent nodes
        if self._is_disfluent_node(node):
            logger.debug(f"Skipping disfluent node: {node.label}")
            return ""
            
        # Process children and join their text
        parts = []
        for child in node.children:
            child_text = self._extract_clean_text(child)
            if child_text:
                parts.append(child_text)
                
        result = " ".join(parts)
        logger.debug(f"Extracted text from {node.label}: '{result}'")
        return result
    
    def _process_speaker_sections(self, text: str) -> Generator[Tuple[str, str], None, None]:
        """
        Split text into speaker sections and yield (speaker_tag, content) pairs.
        Handles both [Speaker] tags and regular text.
        """
        logger.debug(f"Processing speaker sections in text of length: {len(text)}")
        
        # Pattern to match [Speaker] Name patterns
        speaker_pattern = r'(\[Interview(?:er|ee)\][^\n\]]+)(?:\n|$)([\s\S]*?)(?=\n\[|\Z)'
        
        # Find all speaker sections
        sections = list(re.finditer(speaker_pattern, text, re.MULTILINE))
        logger.debug(f"Found {len(sections)} speaker sections")
        
        if not sections:
            # No speaker tags found, yield as a single section with empty speaker
            logger.debug("No speaker tags found, treating as single section")
            yield "", text.strip()
            return
            
        # Process each speaker section
        for i, match in enumerate(sections):
            speaker = match.group(1).strip()
            content = match.group(2).strip()
            logger.debug(f"Processing section {i+1}: speaker='{speaker}', content length={len(content)}")
            
            # Handle text before the first speaker tag
            if i == 0 and match.start() > 0:
                before_text = text[:match.start()].strip()
                if before_text:
                    logger.debug(f"Yielding text before first speaker: '{before_text[:50]}...'")
                    yield "", before_text
            
            # Split content into sentences for better processing
            sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', content) if s.strip()]
            if sentences:
                logger.debug(f"Yielding speaker section with {len(sentences)} sentences")
                yield speaker, ' '.join(sentences)
            
            # Handle text after the last speaker tag
            if i == len(sections) - 1 and match.end() < len(text):
                after_text = text[match.end():].strip()
                if after_text:
                    logger.debug(f"Found text after last speaker: '{after_text[:50]}...'")
                    # Check if there are more speaker tags in the remaining text
                    next_speaker = re.search(r'\[Interview(?:er|ee)\]', after_text)
                    if next_speaker:
                        remaining = after_text[:next_speaker.start()].strip()
                        if remaining:
                            logger.debug(f"Yielding text before next speaker: '{remaining[:50]}...'")
                            yield "", remaining
                    else:
                        logger.debug(f"Yielding final text section: '{after_text[:50]}...'")
                        yield "", after_text
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences while preserving speaker tags and formatting.
        This improved version groups related content to better handle disfluencies."""
        if not text.strip():
            return []
            
        # Clean up the text first
        text = text.strip()
        
        # Split by periods, but be smarter about it
        # We want to group related content together for better disfluency detection
        sentences = []
        
        # Split by periods, but don't split on abbreviations
        parts = re.split(r'(?<=[.!?])\s+', text)
        
        current_sentence_parts = []
        for part in parts:
            if not part.strip():
                continue
                
            # Check if this looks like the start of a new sentence
            # (capitalized and not an abbreviation)
            if (current_sentence_parts and 
                len(part) > 1 and 
                part[0].isupper() and 
                not any(part.lower().startswith(abbr) for abbr in ['mr.', 'mrs.', 'ms.', 'dr.', 'prof.', 'inc.', 'ltd.', 'co.', 'corp.'])):
                
                # Join current sentence parts
                if current_sentence_parts:
                    sentence = ' '.join(current_sentence_parts).strip()
                    if sentence:
                        sentences.append(sentence)
                    current_sentence_parts = []
                    
            current_sentence_parts.append(part)
        
        # Add the last sentence
        if current_sentence_parts:
            sentence = ' '.join(current_sentence_parts).strip()
            if sentence:
                sentences.append(sentence)
        
        # If we have very short sentences, try to combine them for better disfluency detection
        if len(sentences) > 1:
            combined_sentences = []
            current_combined = []
            
            for sentence in sentences:
                words = sentence.split()
                
                # If sentence is very short (less than 5 words), consider combining it
                if len(words) < 5 and current_combined:
                    current_combined.append(sentence)
                else:
                    # If we have accumulated sentences, join them
                    if current_combined:
                        combined_sentences.append(' '.join(current_combined))
                        current_combined = []
                    current_combined = [sentence]
            
            # Add any remaining combined sentences
            if current_combined:
                combined_sentences.append(' '.join(current_combined))
            
            # Only use combined sentences if we actually combined some
            if len(combined_sentences) < len(sentences):
                sentences = combined_sentences
        
        return sentences
    
    def _process_text_section(self, text: str) -> str:
        """Process a single text section to remove disfluencies.
        
        Args:
            text: Text to process
            
        Returns:
            Cleaned text with disfluencies removed
        """
        if not text.strip():
            return ""
        
        # Process the text directly without chunking for better disfluency detection
        # The model works better with complete sentences/paragraphs
        clean_text = self._process_single_chunk(text)
        
        return clean_text
    
    def _process_large_chunk(self, text: str) -> str:
        """Process a large chunk of text to remove disfluencies."""
        if not text.strip():
            return ""
            
        try:
            logger.debug(f"Processing large chunk: '{text[:100]}...'")
            
            # Tokenize the text
            words = text.split()
            if not words:
                return ""
                
            # Use smaller chunks for better disfluency detection
            # The model works better with more context but not too much
            max_words_per_chunk = 30  # Smaller chunks for better disfluency detection
            
            if len(words) <= max_words_per_chunk:
                # Process the entire text as one chunk
                return self._process_single_chunk(text)
            else:
                # Split into chunks and process each
                chunks = []
                for i in range(0, len(words), max_words_per_chunk):
                    chunk_words = words[i:i + max_words_per_chunk]
                    chunk_text = " ".join(chunk_words)
                    processed_chunk = self._process_single_chunk(chunk_text)
                    if processed_chunk:
                        chunks.append(processed_chunk)
                
                return " ".join(chunks)
                
        except Exception as e:
            logger.error(f"Error processing large chunk: {e}", exc_info=True)
            return text
    
    def _process_single_chunk(self, text: str) -> str:
        """Process a single chunk of text through the disfluency model."""
        if not text.strip():
            return ""
            
        try:
            logger.debug(f"Processing chunk: '{text[:100]}...'")
            
            # Tokenize the chunk
            words = text.split()
            if not words:
                return ""
                
            # Prepare input with dummy tags
            tagged_sentence = [(self.dummy_tag, word) for word in words]
            
            # Process through the model
            with torch.no_grad():
                parsed_trees, _ = self.parser.parse_batch([tagged_sentence])
            
            if not parsed_trees or parsed_trees[0] is None:
                logger.warning(f"No parse tree generated for chunk: '{text[:50]}...'")
                return text
                
            treebank_tree = parsed_trees[0].convert()
            
            # Use the simple extraction method that works well
            clean_text = self._extract_clean_text(treebank_tree).strip()
            
            if clean_text:
                # Apply basic punctuation preservation
                result = self._preserve_basic_punctuation(text, clean_text)
                logger.debug(f"Processed chunk. Original: '{text[:50]}...' -> Clean: '{result[:50]}...'")
                return result
                
            logger.debug(f"No clean text extracted from chunk: '{text[:50]}...'")
            return ""
            
        except Exception as e:
            logger.error(f"Error processing chunk '{text[:50]}...': {e}", exc_info=True)
            return text
    
    def _preserve_basic_punctuation(self, original_text: str, clean_text: str) -> str:
        """Apply basic punctuation preservation to clean text."""
        # Split both texts into words
        original_words = original_text.split()
        clean_words = clean_text.split()
        
        # Create a mapping from clean words to original words with punctuation
        word_mapping = {}
        clean_index = 0
        
        for orig_word in original_words:
            # Clean the original word for comparison
            clean_orig = re.sub(r'[^\w]', '', orig_word.lower())
            if clean_orig and clean_index < len(clean_words):
                clean_compare = re.sub(r'[^\w]', '', clean_words[clean_index].lower())
                if clean_orig == clean_compare:
                    word_mapping[clean_words[clean_index]] = orig_word
                    clean_index += 1
        
        # Reconstruct the text with original punctuation
        result_words = []
        for clean_word in clean_words:
            if clean_word in word_mapping:
                result_words.append(word_mapping[clean_word])
            else:
                result_words.append(clean_word)
        
        return ' '.join(result_words)
    
    def _extract_clean_text_with_punctuation(self, original_text: str, tree: InternalTreebankNode) -> str:
        """Extract clean text while preserving punctuation from the original text."""
        # Get the list of words that should be kept (non-disfluent)
        clean_words = self._get_clean_words(tree)
        
        if not clean_words:
            return ""
        
        # Create a mapping of word positions to determine which words to keep
        original_words = original_text.split()
        word_positions = []
        
        # Find the positions of clean words in the original text
        clean_word_index = 0
        for i, word in enumerate(original_words):
            # Clean the word for comparison (remove punctuation)
            clean_word = re.sub(r'[^\w]', '', word.lower())
            if clean_word and clean_word_index < len(clean_words):
                clean_compare = re.sub(r'[^\w]', '', clean_words[clean_word_index].lower())
                if clean_word == clean_compare:
                    word_positions.append(i)
                    clean_word_index += 1
        
        # Reconstruct the text with preserved punctuation
        result_parts = []
        for i, word in enumerate(original_words):
            if i in word_positions:
                result_parts.append(word)
            elif i > 0 and i-1 in word_positions:
                # Keep punctuation that follows clean words
                result_parts.append(word)
        
        return ' '.join(result_parts).strip()
    
    def _get_clean_words(self, node: InternalTreebankNode) -> List[str]:
        """Get the list of clean words (non-disfluent) from the tree."""
        if isinstance(node, LeafTreebankNode):
            return [node.word]
            
        # Skip disfluent nodes
        if self._is_disfluent_node(node):
            logger.debug(f"Skipping disfluent node: {node.label}")
            return []
            
        # Process children and collect their words
        words = []
        for child in node.children:
            child_words = self._get_clean_words(child)
            words.extend(child_words)
            
        return words
    
    def _process_single_sentence(self, sentence: str) -> str:
        """Process a single sentence through the disfluency model."""
        if not sentence.strip():
            logger.debug("Empty sentence, skipping")
            return ""
            
        try:
            logger.debug(f"Processing sentence: '{sentence}'")
            
            # Tokenize the sentence
            words = sentence.split()
            if not words:
                logger.debug("No words in sentence after splitting")
                return ""
                
            # Prepare input with dummy tags
            tagged_sentence = [(self.dummy_tag, word) for word in words]
            logger.debug(f"Tagged sentence with {len(tagged_sentence)} words")
            
            # Process as a single sentence
            with torch.no_grad():
                parsed_trees, _ = self.parser.parse_batch([tagged_sentence])
            
            if not parsed_trees or parsed_trees[0] is None:
                logger.warning(f"No parse tree generated for sentence: '{sentence}'")
                return sentence
                
            treebank_tree = parsed_trees[0].convert()
            
            # Debug: Print the tree structure for short sentences
            if len(words) < 20:
                logger.debug(f"Tree structure for sentence: {treebank_tree}")
            
            clean_text = self._extract_clean_text(treebank_tree).strip()
            
            if clean_text:
                # Ensure proper capitalization
                if clean_text and clean_text[0].isalpha():
                    result = clean_text[0].upper() + clean_text[1:]
                else:
                    result = clean_text
                logger.debug(f"Processed sentence. Original: '{sentence}' -> Clean: '{result}'")
                return result
                
            logger.debug(f"No clean text extracted from sentence: '{sentence}'")
            return ""
            
        except Exception as e:
            logger.error(f"Error processing sentence '{sentence[:50]}...': {e}", exc_info=True)
            return sentence
    
    def _process_sentence_chunk(self, sentences: List[str]) -> List[str]:
        """Process a chunk of sentences through the disfluency model."""
        if not sentences:
            return []
            
        # Process each sentence individually to avoid tensor size issues
        clean_sentences = []
        for sentence in sentences:
            processed = self._process_single_sentence(sentence)
            if processed:
                clean_sentences.append(processed)
                
        return clean_sentences
    
    def process_text(self, text: str) -> str:
        """
        Process a text string to remove disfluencies and return clean, fluent text.
        Preserves speaker tags and handles empty sections.
        
        Args:
            text: Input text to process
            
        Returns:
            Clean, fluent text with disfluencies removed and speaker tags preserved
        """
        logger.info("Starting text processing")
        try:
            # Process each speaker section separately
            sections = list(self._process_speaker_sections(text))
            total_sections = len(sections)
            logger.info(f"Processing {total_sections} speaker sections")
            
            # Calculate progress increments (5% of total sections)
            progress_increment = max(1, total_sections // 20)  # At least 1 section per increment
            next_progress = progress_increment
            
            # Process each section and collect results
            processed_sections = []
            for i, (speaker, content) in enumerate(sections, 1):
                # Log progress in 5% increments
                if i >= next_progress:
                    progress_percent = min(100, (i * 100) // total_sections)
                    print(f"Progress: {progress_percent}% complete ({i}/{total_sections} sections)")
                    next_progress = i + progress_increment
                
                if not content.strip():
                    logger.debug("Empty content, skipping")
                    if speaker:  # Keep speaker tag even if no content
                        processed_sections.append(speaker)
                    continue
                    
                try:
                    # Process the content in smaller chunks
                    logger.debug(f"Processing content: '{content[:100]}...")
                    processed_content = self._process_text_section(content)
                    
                    # Always include the section, but prefer processed content
                    if speaker:
                        processed_sections.append(f"{speaker}\n{processed_content}")
                    else:
                        processed_sections.append(processed_content)
                        
                except Exception as e:
                    logger.error(f"Error processing section: {e}", exc_info=True)
                    # Include the original content if processing fails
                    if speaker:
                        processed_sections.append(f"{speaker}\n{content}")
                    else:
                        processed_sections.append(content)
            
            print("Progress: 100% complete")
            logger.info(f"Processing complete. Processed {len(processed_sections)} sections")
            
            # Join all sections with double newlines
            final_result = "\n\n".join(section for section in processed_sections if section.strip())
            
            # If we somehow ended up with no content, return the original text
            if not final_result.strip():
                logger.warning("No content generated, returning original text")
                return text
                
            return final_result
            
        except Exception as e:
            logger.error(f"Error in process_text: {e}", exc_info=True)
            # Return original text if processing fails completely
            return text

def read_text_file(file_path: Path) -> str:
    """Read text from a .txt file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read().strip()

def read_docx_file(file_path: Path) -> str:
    """Read text from a .docx file."""
    doc = Document(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

def write_text_file(file_path: Path, text: str) -> None:
    """Write text to a .txt file."""
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(text)

def write_docx_file(file_path: Path, text: str) -> None:
    """Write text to a .docx file."""
    doc = Document()
    for paragraph in text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(file_path)

def get_file_type(file_path: Path) -> str:
    """Determine if the file is .txt or .docx based on extension."""
    ext = file_path.suffix.lower()
    if ext == '.txt':
        return 'text'
    elif ext == '.docx':
        return 'docx'
    else:
        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type == 'text/plain':
            return 'text'
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return 'docx'
    raise ValueError(f"Unsupported file type: {file_path}")

def main():
    # Paths
    base_dir = Path(__file__).parent.parent.parent
    model_path = base_dir / 'joint-disfluency-detector-and-parser' / 'best_models' / 'swbd_fisher_bert_Edev.0.9078.pt'
    
    # Get input and output file paths from command line arguments or use defaults
    if len(sys.argv) > 1:
        input_path = Path(sys.argv[1]).resolve()  # Convert to absolute path
    else:
        input_path = base_dir / 'data' / 'raw' / 'input.txt'
    
    if len(sys.argv) > 2:
        output_path = Path(sys.argv[2]).resolve()  # Convert to absolute path
    else:
        # Always use .docx as the default output format
        output_path = input_path.with_stem(f"{input_path.stem}.cleaned").with_suffix('.docx')
    
    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Initialize processor
    processor = DisfluencyProcessor(str(model_path))
    
    try:
        # Read input file based on type
        input_type = get_file_type(input_path)
        print(f"Reading {input_type.upper()} file: {input_path}")
        
        if input_type == 'text':
            text = read_text_file(input_path)
        else:  # docx
            text = read_docx_file(input_path)
        
        if not text.strip():
            print("Error: Input file is empty.")
            return
        
        # Process text to get clean output
        print("Processing text and removing disfluencies...")
        clean_text = processor.process_text(text)
        
        # Write output based on output file extension
        output_type = get_file_type(output_path)
        print(f"Writing {output_type.upper()} file: {output_path}")
        
        if output_type == 'text':
            write_text_file(output_path, clean_text)
        else:  # docx
            write_docx_file(output_path, clean_text)
        
        print(f"\nProcessing complete. Clean text written to {output_path}")
        
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
